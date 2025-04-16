from docx import Document
from chathrd.parsers.utils import get_file_title, logger

def parse_docx(file_path: str) -> str:
    """Парсит DOCX файл и возвращает его содержимое в markdown формате."""
    try:
        doc = Document(file_path)
        content = [f"# {get_file_title(file_path)}", ""]

        # Обработка текста
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        if text_content:
            content.extend(["## Текст документа", ""] + text_content + [""])

        # Обработка таблиц
        tables_content = []
        for i, table in enumerate(doc.tables, 1):
            if not any(cell.text.strip() for row in table.rows for cell in row.cells):
                continue

            table_content = [f"### Таблица {i}", ""]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            table_content.append("| " + " | ".join(headers) + " |")
            table_content.append("| " + " | ".join(["---"] * len(headers)) + " |")

            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    table_content.append("| " + " | ".join(cells) + " |")

            tables_content.extend(table_content + [""])

        if tables_content:
            content.extend(["### Извлеченные таблицы:", ""] + tables_content)

        return '\n'.join(content)
    except Exception as e:
        logger.error(f"Ошибка при парсинге DOCX файла {file_path}: {str(e)}")
        return f"Ошибка при обработке DOCX файла: {str(e)}" 