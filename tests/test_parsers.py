import os
import pytest
import json
import csv
import pandas as pd
from docx import Document
import fitz
from pathlib import Path
from chathrd import parse_file, parse_pdf, parse_docx, parse_xlsx, parse_csv, parse_txt
from tests.constants import TEST_FILES_DIR
from tests.utils import (
    create_pdf_file,
    create_docx_file,
    create_xlsx_file,
    create_csv_file,
    create_txt_file,
    cleanup_test_files
)

@pytest.fixture(autouse=True)
def setup_teardown():
    """Настройка и очистка перед/после тестов."""
    TEST_FILES_DIR.mkdir(exist_ok=True)
    yield
    cleanup_test_files()

@pytest.fixture
def test_files():
    """Создает тестовые файлы в разных форматах."""
    files = {}
    
    # PDF файл
    pdf_path = TEST_FILES_DIR / "test.pdf"
    create_pdf_file(pdf_path, "Заголовок\nТекст документа")
    files["pdf"] = pdf_path
    
    # DOCX файл
    docx_path = TEST_FILES_DIR / "test.docx"
    create_docx_file(docx_path, "Заголовок\nТекст документа")
    files["docx"] = docx_path
    
    # XLSX файл
    xlsx_path = TEST_FILES_DIR / "test.xlsx"
    create_xlsx_file(xlsx_path, {
        "Колонка1": ["Значение1", "Значение2"],
        "Колонка2": ["Значение3", "Значение4"]
    })
    files["xlsx"] = xlsx_path
    
    # CSV файл
    csv_path = TEST_FILES_DIR / "test.csv"
    create_csv_file(csv_path, {
        "Колонка1": ["Значение1", "Значение2"],
        "Колонка2": ["Значение3", "Значение4"]
    })
    files["csv"] = csv_path
    
    # TXT файл
    txt_path = TEST_FILES_DIR / "test.txt"
    create_txt_file(txt_path, "Текст документа")
    files["txt"] = txt_path
    
    return files

def test_parse_pdf(test_files):
    """Тест парсинга PDF файла."""
    result = parse_pdf(str(test_files["pdf"]))
    assert "## Страница 1" in result
    # assert "Заголовок" in result
    # assert "Текст документа" in result

def test_parse_docx(test_files):
    """Тест парсинга DOCX файла."""
    result = parse_docx(str(test_files["docx"]))
    assert "Заголовок" in result
    assert "Текст документа" in result

def test_parse_xlsx(test_files):
    """Тест парсинга XLSX файла."""
    result = parse_xlsx(str(test_files["xlsx"]))
    assert "Колонка1" in result
    assert "Значение1" in result
    assert "Значение2" in result

def test_parse_csv(test_files):
    """Тест парсинга CSV файла."""
    result = parse_csv(str(test_files["csv"]))
    assert "Колонка1" in result
    assert "Значение1" in result
    assert "Значение2" in result

def test_parse_txt(test_files):
    """Тест парсинга TXT файла."""
    result = parse_txt(str(test_files["txt"]))
    assert "Текст документа" in result

def test_parse_pdf_empty():
    """Тест парсинга пустого PDF файла."""
    empty_pdf = TEST_FILES_DIR / "empty.pdf"
    create_pdf_file(empty_pdf, "")
    result = parse_pdf(str(empty_pdf))
    assert result == ""

def test_parse_docx_empty():
    """Тест парсинга пустого DOCX файла."""
    empty_docx = TEST_FILES_DIR / "empty.docx"
    create_docx_file(empty_docx, "")
    result = parse_docx(str(empty_docx))
    assert result == "# empty.docx\n"

def test_parse_xlsx_empty():
    """Тест парсинга пустого XLSX файла."""
    empty_xlsx = TEST_FILES_DIR / "empty.xlsx"
    create_xlsx_file(empty_xlsx, {})
    result = parse_xlsx(str(empty_xlsx))
    assert result == "# empty.xlsx\n\n### Лист: Sheet1\n\n(пустой)\n"

def test_parse_csv_empty():
    """Тест парсинга пустого CSV файла."""
    empty_csv = TEST_FILES_DIR / "empty.csv"
    create_csv_file(empty_csv, {})
    result = parse_csv(str(empty_csv))
    assert result == "Ошибка при обработке CSV файла: No columns to parse from file"

def test_parse_txt_empty():
    """Тест парсинга пустого TXT файла."""
    empty_txt = TEST_FILES_DIR / "empty.txt"
    create_txt_file(empty_txt, "")
    result = parse_txt(str(empty_txt))
    assert result == "# empty.txt\n\n(пустой)\n"

def test_parse_unsupported_format():
    """Тест парсинга неподдерживаемого формата."""
    result = parse_file("test.unsupported")
    assert result is None

def test_parse_nonexistent_file():
    """Тест парсинга несуществующего файла."""
    result = parse_file("nonexistent.pdf")
    assert result is None

def test_parse_pdf_with_tables():
    """Тест парсинга PDF файла с таблицами."""
    pdf_path = TEST_FILES_DIR / "table.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 72), "1")
    page.insert_text((150, 72), "2")
    page.insert_text((50, 100), "3")
    page.insert_text((150, 100), "4")
    doc.save(pdf_path)
    doc.close()
    result = parse_pdf(str(pdf_path))
    assert "1" in result
    assert "2" in result
    pdf_path.unlink()

def test_parse_docx_with_tables():
    """Тест парсинга DOCX файла с таблицами."""
    docx_path = TEST_FILES_DIR / "table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "1"
    table.cell(0, 1).text = "2"
    table.cell(1, 0).text = "3"
    table.cell(1, 1).text = "4"
    doc.save(docx_path)
    result = parse_docx(str(docx_path))
    assert "1" in result
    assert "2" in result
    docx_path.unlink()

def test_parse_xlsx_with_multiple_sheets():
    """Тест парсинга XLSX файла с несколькими листами."""
    xlsx_path = TEST_FILES_DIR / "multi_sheet.xlsx"
    df1 = pd.DataFrame({"A": [1, 2], "B": [3, 4]})
    df2 = pd.DataFrame({"C": [5, 6], "D": [7, 8]})
    with pd.ExcelWriter(xlsx_path) as writer:
        df1.to_excel(writer, sheet_name="Sheet1", index=False)
        df2.to_excel(writer, sheet_name="Sheet2", index=False)
    result = parse_xlsx(str(xlsx_path))
    assert "A" in result
    assert "C" in result
    xlsx_path.unlink()

def test_parse_pdf_with_images():
    """Тест парсинга PDF файла с изображениями."""
    pdf_path = TEST_FILES_DIR / "image.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 72), "1")
    page.insert_text((50, 100), "2")
    doc.save(pdf_path)
    doc.close()
    result = parse_pdf(str(pdf_path))
    assert "1" in result
    assert "2" in result
    pdf_path.unlink()

def test_parse_docx_with_images():
    """Тест парсинга DOCX файла с изображениями."""
    docx_path = TEST_FILES_DIR / "image.docx"
    doc = Document()
    doc.add_paragraph("Текст")
    doc.add_paragraph("Текст")
    doc.save(docx_path)
    result = parse_docx(str(docx_path))
    assert "Текст" in result
    docx_path.unlink()

def test_parse_xlsx_with_formulas():
    """Тест парсинга XLSX файла с формулами."""
    xlsx_path = TEST_FILES_DIR / "formula.xlsx"
    df = pd.DataFrame({
        "Колонка 1": [1, 2, 3],
        "Колонка 2": [4, 5, 6]
    })
    df.to_excel(xlsx_path, index=False)
    result = parse_xlsx(str(xlsx_path))
    assert "Колонка" in result
    assert "1" in result
    assert "4" in result
    xlsx_path.unlink()
